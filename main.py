# -*- coding: utf-8 -*-
# 5.4更新：添加了日志记录模块，优化了错误处理逻辑，现在的逻辑是一遍过：出现错误就处理，处理三次还解决不了的话就跳过
import os

import docx
import openai
import logging
import sys
import requests
import wget
import time, hashlib
from qiniu import Auth, put_file, etag

access_key = '_kE1kDRGOVlOrEjhd8Y2gUYzyvoy0od274RTC8Py'
secret_key = '4s-v5JfrhLf5c6WCQ507OMHd-_LLIKtZ0LjFEg3I'
openai_api_key = "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"

# 移动p1段落到p2段落后面
def move_paragraph_after(para1, para2):
    """移动p1段落到p2段落后面"""
    p1, p2 = para1._p, para2._p
    p2.addnext(p1)

# 在指定的段落pa后面插入段落
def insert_paragraph_after(paragraph, text=None):
    """在指定的段落后面创建段落"""
    sy = paragraph.style
    # print(len(paragraph.runs))
    para = paragraph.insert_paragraph_before(text, style=sy)
    for run in para.runs:
        run.font.color.rgb = docx.shared.RGBColor(34, 139, 34)
    move_paragraph_after(para, paragraph)
    return para

# 调用API
# 通过是否使用第三方服务控制host,api_bash
# openai.api_base = 'https://api.closeai-asia.com/v1'
def getResult(rw:str,flag:bool="True",key:str="",prompt:str="")->str:
    # 使用第三方服务
    if flag:
        sleeptime=0
        openai.api_base = 'https://api.closeai-asia.com/v1'
        openai.api_key = key
    # openai.host=host
    else:
        # 使用OPENAI服务
        proxy = {
            'http': 'http://localhost:7890',
            'https': 'http://localhost:7890'
        }

        openai.api_key = "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"
        openai.proxy = proxy
        sleeptime=20
    time.sleep(sleeptime)
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": rw},
        ]
    )
    result = response['choices'][0]['message']['content'].strip()
    result = result.replace('\n', '').strip().replace('\r', '')
    print(result)
    return result

# 检查这一段需不需要调用API
def check_Necessity(para, select_mode="only_red"):
    if len(para.text) < 23:
        # insert_paragraph_after(para, text="这一段字数少于20字 %d\t\n" %(count,total) +)
        return False
    elif select_mode == "only_red":
        for run in para.runs:
            # 如果字体颜色为红色
            if run.font.color.rgb == docx.shared.RGBColor(255, 0, 0):
                return True
        return False
    return True


# 对有问题的段落再进行处理
def deal_error_para(rw='',flag:bool=True,key='',prompt=''):
    count = 0
    while count <= 3:
        try:
            time.sleep(20)
            result = getResult(flag=flag,key=key,prompt=prompt,rw=rw.text)
            logging.info("This error has dealed")
            return result
        except Exception as e:
            count += 1
            logging.error("Deal with error para\n" + str(e) + "【原文】" + rw.text + "\n【ERROR】" + str(e) + "\n")
    logging.error("这段处理不了了")
    return "这段处理不了"

# fastapi创建的时候创建好可以复用的参数
def download_file(host="rv548d2bf.hb-bkt.clouddn.com",file_key="0524-1.docx"):
    q = Auth(access_key, secret_key)

    # 有两种方式构造base_url的形式
    base_url = 'http://%s/%s' % (host, file_key)

    # 可以设置token过期时间
    private_url = q.private_download_url(base_url, expires=3600)

    # 判断链接是否有效，这里应该让抛出的异常传至控制模块
    r = requests.get(private_url)
    assert r.status_code == 200

    file_position = "./download"
    # 保证文件夹一定存在
    if not os.path.exists(file_position):
        os.makedirs(file_position)
    # 构造下载后的文件路径然后下载，之后返回该路径
    file_path = file_position + "/%s" % file_key
    try:
        wget.download(private_url, file_path)
    except Exception as e:
        print("文件无法下载")
        return{"result":0,"info":"文件无法下载"}
    print("下载完成，位置：", file_path)
    return dict(result=1,info=file_path)


def upload_file(localfile, bucket_name='zxzapi0524', key="test_result.docx"):
    # 构建鉴权对象
    q = Auth(access_key, secret_key)
    # 生成上传 Token，可以指定过期时间等
    token = q.upload_token(bucket_name, key, 3600)
    try:
        ret, info = put_file(token, key, localfile, version='v2')
        print(ret,info)
    except Exception as e:
        print("file can't upload")
        print(e)
        info="file can't upload"
    try:
        os.remove(localfile)
        print("结果文件已删除")
    except Exception as e:
        print("结果文件无法删除")
    print(info)
    # assert ret['key'] == key
    # assert ret['hash'] == etag(file)
    return {"result_code":info,"result_key":key}


# files=[filepath]
# 输出到控制台
"""
select_mode : only_red,all
write_mode : insert,replace
"""

def processing(file, flag,key,prompt="重写：", select_mode="only_red", write_mode="insert"):
    LOG_FORMAT = "%(asctime)s - %(levelname)s - %(message)s"
    logging.basicConfig(filename='%s.log' % time.strftime("%m%d"), level=logging.INFO, format=LOG_FORMAT)
    # 打开docx文档
    doc = docx.Document(file)
    count=0
    total=len(doc.paragraphs)
    for para in doc.paragraphs:
        count += 1
        if check_Necessity(para, select_mode):
            try:
                print("NO.%d/%d start -f:%s" % (count, total, file))
                if write_mode == "replace":
                    # 这个是替换
                    para.text = getResult(flag=flag,key=key,prompt=prompt,rw=para.text)
                    for run in para.runs:
                        run.font.color.rgb = docx.shared.RGBColor(34, 139, 34)
                else:
                    # 下面这个是插入
                    insert_paragraph_after(para, getResult(flag=flag,key=key,prompt=prompt,rw=para.text))
            except Exception as e:
                print("NO.%d/%d has Error -f:%s" % (count, total, file) + str(e))
                if write_mode == "replace":
                    # 这个是替换
                    para.text = deal_error_para(flag=flag,key=key,prompt=prompt,rw=para.text)
                    for run in para.runs:
                        run.font.color.rgb = docx.shared.RGBColor(34, 139, 34)
                else:
                    # 下面这个是插入
                    insert_paragraph_after(para, deal_error_para(flag=flag,key=key,prompt=prompt,rw=para.text))
                    print("NO.%d/%d has Error -f :%s" % (count, total, file) + str(e))
            print("NO.%d/%d done -f:%s" % (count, total, file))
        else:
            print("NO.%d/%d para skipped -f:%s" % (count, total, file))
    try:
        # 这里还要改
        file_path = file.replace(".docx", "") + "_result.docx"
        doc.save(file_path)
        print("程序已运行完毕，文件路径：",file_path)
    except Exception as e:
        logging.error("FILE : 【%s】 can't save." % file)
        print("error:file can't save", e)
        file_path = file.replace(".docx", ".txt")
        with open(file, "w+") as f:
            for i in doc.paragraphs:
                f.write(i.text)
        print("txt file saved in ", file)
    return file_path
if __name__ == '__main__':
    """
    download_file的两个参数都是通过命令行传入的，第一个是host域名。第二个是key
    这里download_file()应该由命令行传入参数，通过fastapi建立的接口传入参数
    可以构建一个微服务,传入相关参数，然后写一个主控逻辑控制任务进行
    暴露的外部接口是【七牛云的空间】、【七牛云的key】、【openai_api_key】、【完成任务需要调用的链接】,只需要用fastapi接收这两个参数然后调用各个组件就行
    发送一个post请求，附带json格式的数据：输出的是一个json格式的数据，包含【任务是否成功】、【失败原因】、【七牛云的空间】、【七牛云的key】
    
    """
    # 在线处理逻辑
    '''
    filepath = download_file("rv548d2bf.hb-bkt.clouddn.com", "test.docx")
    print("参数是：\t", sys.argv)
    # source_files=walkFile(filepath)
    result_file_path = processing([filepath], "only_red", "insert")
    res = upload_file(localfile=result_file_path, bucket_name="zxzapi0524", key="test_res_0524.docx")
    print("result is :", res)
    '''
    # 本地处理逻辑
    # 所有参数
    select_mode = "only_red"
    write_mode = "insert"
    pmt = "我希望你能充当一个认真仔细的改重资深工作人员，你需要将不同方法融合起来将这段文字词和短语进行彻底全面细致改写，具体方法有同义词替换、详细阐述、换角度描述、增加过度语句等，使得与原来的内容不会被认定为抄袭，同时要确保语句通顺，内容符合逻辑"

    file = r"./企业流转土地规模经营失败的原因探究(1).docx"
    result_file_path = processing(file,prompt=pmt, select_mode=select_mode, write_mode=write_mode)


